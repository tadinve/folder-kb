"""
Document Processing Pipeline for Construction Knowledge Graph

This module handles parsing various document types and extracting entities
for knowledge graph population.
"""

import os
import json
import pandas as pd
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from datetime import datetime

# Document processing
import PyPDF2
from docx import Document
import openpyxl

# NLP and ML
import spacy
from sentence_transformers import SentenceTransformer
import numpy as np

# Knowledge graph
from neo4j import GraphDatabase

@dataclass
class ProcessedDocument:
    """Container for processed document information."""
    file_path: str
    filename: str
    file_type: str
    content: str
    entities: Dict[str, List[str]]
    relationships: List[Dict[str, Any]]
    embedding: Optional[np.ndarray] = None
    metadata: Optional[Dict[str, Any]] = None

class DocumentProcessor:
    """Processes construction project documents for knowledge graph ingestion."""
    
    def __init__(self, embedding_model: str = "all-MiniLM-L6-v2"):
        """Initialize the document processor."""
        self.embedding_model = SentenceTransformer(embedding_model)
        
        # Load spaCy model for entity extraction
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            print("⚠️  spaCy model not found. Install with: python -m spacy download en_core_web_sm")
            self.nlp = None
        
        # Construction-specific entity patterns
        self.construction_patterns = {
            'project_codes': r'(PCO|RFI|SI|CO)\s*[#-]?\s*\d+',
            'drawing_numbers': r'[A-Z]-[A-Z0-9]+-\d+',
            'dates': r'\d{4}[-/.]\d{1,2}[-/.]\d{1,2}',
            'costs': r'\$[\d,]+\.?\d*',
            'phases': r'Phase\s+\d+|Area\s+\d+|Building\s+\d+',
        }
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            print(f"❌ Error processing PDF {file_path}: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from Word document."""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            print(f"❌ Error processing DOCX {file_path}: {e}")
            return ""
    
    def extract_text_from_excel(self, file_path: str) -> str:
        """Extract text from Excel file."""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text = ""
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"Sheet: {sheet_name}\n"
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        text += row_text + "\n"
                text += "\n"
            
            return text.strip()
        except Exception as e:
            print(f"❌ Error processing Excel {file_path}: {e}")
            return ""
    
    def extract_entities(self, text: str) -> Dict[str, List[str]]:
        """Extract entities from text using spaCy and custom patterns."""
        entities = {
            'persons': [],
            'organizations': [],
            'dates': [],
            'money': [],
            'locations': [],
            'project_codes': [],
            'drawing_numbers': [],
            'phases': []
        }
        
        if not self.nlp:
            return entities
        
        # Process with spaCy
        doc = self.nlp(text)
        
        # Extract named entities
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                entities['persons'].append(ent.text)
            elif ent.label_ == "ORG":
                entities['organizations'].append(ent.text)
            elif ent.label_ in ["DATE", "TIME"]:
                entities['dates'].append(ent.text)
            elif ent.label_ == "MONEY":
                entities['money'].append(ent.text)
            elif ent.label_ in ["GPE", "LOC"]:
                entities['locations'].append(ent.text)
        
        # Extract construction-specific patterns
        import re
        for pattern_name, pattern in self.construction_patterns.items():
            matches = re.findall(pattern, text, re.IGNORECASE)
            if pattern_name in entities:
                entities[pattern_name].extend(matches)
        
        # Remove duplicates and clean
        for key in entities:
            entities[key] = list(set([item.strip() for item in entities[key] if item.strip()]))
        
        return entities
    
    def detect_relationships(self, text: str, entities: Dict[str, List[str]]) -> List[Dict[str, Any]]:
        """Detect relationships between entities in the text."""
        relationships = []
        
        # Simple relationship patterns for construction documents
        relationship_patterns = [
            {
                'pattern': r'(.*?)\s+approved\s+(.*?)(?:\.|$)',
                'relation': 'APPROVED',
                'subject_group': 1,
                'object_group': 2
            },
            {
                'pattern': r'(.*?)\s+submitted\s+(.*?)(?:\.|$)',
                'relation': 'SUBMITTED',
                'subject_group': 1,
                'object_group': 2
            },
            {
                'pattern': r'(.*?)\s+responsible\s+for\s+(.*?)(?:\.|$)',
                'relation': 'RESPONSIBLE_FOR',
                'subject_group': 1,
                'object_group': 2
            },
            {
                'pattern': r'(.*?)\s+assigned\s+to\s+(.*?)(?:\.|$)',
                'relation': 'ASSIGNED_TO',
                'subject_group': 1,
                'object_group': 2
            }
        ]
        
        import re
        for pattern_info in relationship_patterns:
            matches = re.finditer(pattern_info['pattern'], text, re.IGNORECASE)
            for match in matches:
                subject = match.group(pattern_info['subject_group']).strip()
                obj = match.group(pattern_info['object_group']).strip()
                
                if len(subject) > 3 and len(obj) > 3:  # Filter out very short matches
                    relationships.append({
                        'subject': subject,
                        'relation': pattern_info['relation'],
                        'object': obj,
                        'context': match.group(0)
                    })
        
        return relationships
    
    def generate_embedding(self, text: str) -> np.ndarray:
        """Generate embedding for the document text."""
        # Chunk text if it's too long
        max_length = 500
        if len(text) > max_length:
            # Take first part of the document for embedding
            text = text[:max_length]
        
        embedding = self.embedding_model.encode(text)
        return embedding
    
    def process_document(self, file_path: str) -> Optional[ProcessedDocument]:
        """Process a single document and extract all information."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            print(f"❌ File not found: {file_path}")
            return None
        
        print(f"📄 Processing: {file_path.name}")
        
        # Extract text based on file type
        file_type = file_path.suffix.lower()
        
        if file_type == '.pdf':
            text = self.extract_text_from_pdf(str(file_path))
        elif file_type == '.docx':
            text = self.extract_text_from_docx(str(file_path))
        elif file_type in ['.xlsx', '.xls']:
            text = self.extract_text_from_excel(str(file_path))
        else:
            print(f"⚠️  Unsupported file type: {file_type}")
            return None
        
        if not text:
            print(f"⚠️  No text extracted from: {file_path.name}")
            return None
        
        # Extract entities and relationships
        entities = self.extract_entities(text)
        relationships = self.detect_relationships(text, entities)
        
        # Generate embedding
        embedding = self.generate_embedding(text)
        
        # Create metadata
        stat_info = file_path.stat()
        metadata = {
            'file_size': stat_info.st_size,
            'created_date': datetime.fromtimestamp(stat_info.st_ctime).isoformat(),
            'modified_date': datetime.fromtimestamp(stat_info.st_mtime).isoformat(),
            'text_length': len(text),
            'entity_count': sum(len(v) for v in entities.values()),
            'relationship_count': len(relationships)
        }
        
        return ProcessedDocument(
            file_path=str(file_path),
            filename=file_path.name,
            file_type=file_type.lstrip('.'),
            content=text,
            entities=entities,
            relationships=relationships,
            embedding=embedding,
            metadata=metadata
        )
    
    def process_directory(self, directory_path: str, file_types: List[str] = None) -> List[ProcessedDocument]:
        """Process all documents in a directory."""
        if file_types is None:
            file_types = ['.pdf', '.docx', '.xlsx', '.xls']
        
        directory_path = Path(directory_path)
        processed_docs = []
        
        print(f"🔍 Scanning directory: {directory_path}")
        
        for file_path in directory_path.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in file_types:
                doc = self.process_document(file_path)
                if doc:
                    processed_docs.append(doc)
        
        print(f"✅ Processed {len(processed_docs)} documents")
        return processed_docs
    
    def save_processed_data(self, processed_docs: List[ProcessedDocument], output_dir: str):
        """Save processed document data to files."""
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Save document metadata
        metadata_list = []
        for doc in processed_docs:
            metadata_list.append({
                'file_path': doc.file_path,
                'filename': doc.filename,
                'file_type': doc.file_type,
                'metadata': doc.metadata,
                'entity_counts': {k: len(v) for k, v in doc.entities.items()},
                'relationship_count': len(doc.relationships)
            })
        
        metadata_df = pd.DataFrame(metadata_list)
        metadata_df.to_csv(output_dir / 'document_metadata.csv', index=False)
        
        # Save full processed data as JSON
        processed_data = []
        for doc in processed_docs:
            doc_data = {
                'file_path': doc.file_path,
                'filename': doc.filename,
                'file_type': doc.file_type,
                'content': doc.content[:1000],  # Truncate for JSON
                'entities': doc.entities,
                'relationships': doc.relationships,
                'metadata': doc.metadata
            }
            processed_data.append(doc_data)
        
        with open(output_dir / 'processed_documents.json', 'w') as f:
            json.dump(processed_data, f, indent=2)
        
        # Save embeddings separately
        embeddings = np.array([doc.embedding for doc in processed_docs])
        np.save(output_dir / 'document_embeddings.npy', embeddings)
        
        print(f"💾 Saved processed data to: {output_dir}")

def main():
    """Example usage of the document processor."""
    
    # Initialize processor
    processor = DocumentProcessor()
    
    # Process documents from the construction project
    input_dir = "/Users/venkateshtadinada/Documents/VS-Code-Projects/folder-kb/DocLabs_Sample_Project_Template"
    output_dir = "/Users/venkateshtadinada/Documents/VS-Code-Projects/folder-kb/data/processed"
    
    print("🚀 Starting document processing pipeline...")
    
    # Process a limited number of files for demonstration
    processed_docs = []
    
    # Get first 10 PDF files for testing
    pdf_files = list(Path(input_dir).rglob('*.pdf'))[:10]
    
    for pdf_file in pdf_files:
        doc = processor.process_document(pdf_file)
        if doc:
            processed_docs.append(doc)
    
    if processed_docs:
        # Save results
        processor.save_processed_data(processed_docs, output_dir)
        
        # Print summary
        print(f"\n📊 Processing Summary:")
        print(f"Documents processed: {len(processed_docs)}")
        print(f"Total entities extracted: {sum(sum(len(v) for v in doc.entities.values()) for doc in processed_docs)}")
        print(f"Total relationships found: {sum(len(doc.relationships) for doc in processed_docs)}")
    else:
        print("❌ No documents were successfully processed")

if __name__ == "__main__":
    main()